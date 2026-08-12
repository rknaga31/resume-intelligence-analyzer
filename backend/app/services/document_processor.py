"""
Document Processing Service — Milestone 3.

Handles:
- MIME type validation
- File header (magic byte) verification
- Size limit enforcement
- PDF text extraction (multi-backend fallback)
- DOCX text extraction
- TXT text extraction
- Empty / corrupted document detection
- Path traversal prevention

SECURITY: Never log extracted text content. Never persist files to disk.
"""
from __future__ import annotations

import io
import os
import re
import uuid

from app.core.config import get_settings
from app.core.exceptions import (
    DocumentProcessingError,
    EmptyDocumentError,
    FileTooLargeError,
    UnsupportedFileTypeError,
)
from app.core.logging import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Magic Bytes
# ---------------------------------------------------------------------------
_MAGIC_BYTES: dict[str, bytes] = {
    "application/pdf": b"%PDF-",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": b"PK\x03\x04",
    "application/msword": b"\xd0\xcf\x11\xe0",  # OLE2 compound doc
    "text/plain": b"",  # No magic bytes for plain text
}

# Minimum extracted text length to be considered a valid resume
_MIN_TEXT_LENGTH = 50


class DocumentProcessor:
    """Validates and extracts text from uploaded resume documents.

    All processing is in-memory — no files are written to disk.
    """

    def __init__(self) -> None:
        """Initialise processor with application settings."""
        self._settings = get_settings()

    def validate_and_extract(
        self,
        file_content: bytes,
        filename: str,
        declared_content_type: str,
    ) -> dict:
        """Validate a resume file and extract its text content.

        Args:
            file_content: Raw bytes of the uploaded file.
            filename: Original filename from the upload (sanitized internally).
            declared_content_type: MIME type declared by the HTTP client.

        Returns:
            Dictionary with extracted text, metadata, and extraction method used.

        Raises:
            FileTooLargeError: File exceeds the configured size limit.
            UnsupportedFileTypeError: File type is not supported.
            DocumentProcessingError: Extraction failed.
            EmptyDocumentError: Extracted text is too short to be a valid resume.
        """
        # 1. Size check
        self._check_size(file_content)

        # 2. Sanitize filename (prevent path traversal)
        safe_filename = self._sanitize_filename(filename)

        # 3. Detect MIME type from magic bytes (do not trust client declaration)
        detected_mime = self._detect_mime(file_content, safe_filename)

        # 4. Validate against allowed types
        if detected_mime not in self._settings.allowed_mime_types:
            raise UnsupportedFileTypeError(
                f"File type '{detected_mime}' is not supported. "
                f"Please upload a PDF, DOCX, or TXT file.",
                details={"detected_mime": detected_mime},
            )

        # 5. Extract text
        text, extraction_method = self._extract_text(file_content, detected_mime)

        # 6. Validate extracted content
        cleaned_text = self._clean_text(text)
        if len(cleaned_text) < _MIN_TEXT_LENGTH:
            raise EmptyDocumentError(
                "The uploaded document appears to be empty or contains insufficient text. "
                "Please ensure your resume is not password-protected or image-only.",
                details={"extracted_length": len(cleaned_text)},
            )

        file_id = f"res_{uuid.uuid4().hex[:8]}"
        logger.info(
            "document_processed",
            file_id=file_id,
            mime_type=detected_mime,
            file_size=len(file_content),
            word_count=len(cleaned_text.split()),
            extraction_method=extraction_method,
            # DO NOT log: safe_filename or any extracted text content
        )

        return {
            "file_id": file_id,
            "filename": safe_filename,
            "file_size": len(file_content),
            "mime_type": detected_mime,
            "extracted_text": cleaned_text,
            "word_count": len(cleaned_text.split()),
            "character_count": len(cleaned_text),
            "extraction_method": extraction_method,
        }

    # -------------------------------------------------------------------------
    # Private helpers
    # -------------------------------------------------------------------------

    def _check_size(self, content: bytes) -> None:
        """Raise FileTooLargeError if content exceeds the configured limit."""
        if len(content) > self._settings.max_upload_size_bytes:
            raise FileTooLargeError(
                f"File size {len(content):,} bytes exceeds the maximum allowed "
                f"size of {self._settings.max_upload_size_mb} MB.",
                details={"max_bytes": self._settings.max_upload_size_bytes},
            )

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """Strip path components and dangerous characters from a filename."""
        basename = os.path.basename(filename)
        # Replace any non-alphanumeric/dash/dot/underscore characters
        safe = re.sub(r"[^\w.\-]", "_", basename)
        return safe[:255] or "resume"

    @staticmethod
    def _detect_mime(content: bytes, filename: str) -> str:
        """Detect MIME type from file header bytes, falling back to extension."""
        # Check magic bytes
        for mime, magic in _MAGIC_BYTES.items():
            if magic and content.startswith(magic):
                return mime

        # Plain text fallback (no magic bytes)
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "txt":
            return "text/plain"
        if ext in ("pdf",):
            return "application/pdf"
        if ext in ("docx",):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext in ("doc",):
            return "application/msword"

        # Return application/octet-stream to trigger unsupported error
        return "application/octet-stream"

    def _extract_text(self, content: bytes, mime_type: str) -> tuple[str, str]:
        """Route extraction to the appropriate backend.

        Args:
            content: Raw file bytes.
            mime_type: Detected MIME type.

        Returns:
            Tuple of (extracted_text, extraction_method_name).
        """
        if mime_type == "application/pdf":
            return self._extract_pdf(content)
        if mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return self._extract_docx(content)
        if mime_type == "text/plain":
            return self._extract_txt(content)
        raise UnsupportedFileTypeError(
            f"No extraction backend for MIME type: {mime_type}"
        )

    @staticmethod
    def _extract_pdf(content: bytes) -> tuple[str, str]:
        """Extract text from PDF bytes using pypdf with pdfplumber fallback."""
        # Primary: pypdf
        try:
            import pypdf  # noqa: PLC0415

            reader = pypdf.PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                raise DocumentProcessingError(
                    "The uploaded PDF is password-protected. "
                    "Please remove the password before uploading."
                )
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
            if text.strip():
                return text, "pypdf"
        except DocumentProcessingError:
            raise
        except Exception:  # noqa: BLE001
            pass  # Fall through to pdfplumber

        # Fallback: pdfplumber
        try:
            import pdfplumber  # noqa: PLC0415

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages)
            if text.strip():
                return text, "pdfplumber"
        except Exception as exc:  # noqa: BLE001
            raise DocumentProcessingError(
                "Failed to extract text from the PDF. "
                "The file may be corrupted, image-only, or in an unsupported format.",
                details={"error_type": type(exc).__name__},
            ) from exc

        raise DocumentProcessingError(
            "PDF appears to contain no extractable text. "
            "If this is a scanned PDF, please convert it to a text-based format."
        )

    @staticmethod
    def _extract_docx(content: bytes) -> tuple[str, str]:
        """Extract text from DOCX bytes using python-docx."""
        try:
            from docx import Document  # noqa: PLC0415

            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            text = "\n".join(p for p in paragraphs if p.strip())
            return text, "python-docx"
        except Exception as exc:  # noqa: BLE001
            raise DocumentProcessingError(
                "Failed to extract text from the DOCX file. "
                "The file may be corrupted or in an unsupported format.",
                details={"error_type": type(exc).__name__},
            ) from exc

    @staticmethod
    def _extract_txt(content: bytes) -> tuple[str, str]:
        """Decode plain-text file, trying common encodings."""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                text = content.decode(encoding)
                return text, f"text-decode-{encoding}"
            except UnicodeDecodeError:
                continue
        raise DocumentProcessingError(
            "Could not decode the text file. Please ensure it is UTF-8 encoded."
        )

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text: normalise whitespace and remove control chars."""
        # Remove null bytes and other control characters except newlines/tabs
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
        # Normalise multiple spaces/blank lines
        text = re.sub(r" {3,}", " ", text)
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        return text.strip()
